import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt

# Корневая директория для поиска
CLIENT_ROOT = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "project_activity_client-master"
)

EXPORT_EXTS = {".js", ".jsx", ".ts", ".tsx", ".css", ".scss"}
EXCLUDE_DIRS = {"node_modules", "venv", ".git", "__pycache__"}
EXCLUDE_BIN_EXTS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".svg",
    ".woff",
    ".woff2",
    ".ttf",
    ".eot",
    ".otf",
    ".ico",
    ".pdf",
}
OUT_FILENAME = "client_sources_code.docx"
CODE_FONT = "Consolas"
CODE_FONT_SIZE = Pt(9)


def is_source_file(filepath: str) -> bool:
    ext = os.path.splitext(filepath)[1].lower()
    if ext in EXPORT_EXTS and ext not in EXCLUDE_BIN_EXTS:
        return True
    return False


def walk_client_files(base_dir: str):
    res = []
    for root, dirs, files in os.walk(base_dir):
        # Исключаем лишние директории
        dirs[:] = [d for d in dirs if d not in EXCLUDE_DIRS]
        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext in EXCLUDE_BIN_EXTS:
                continue
            abs_path = os.path.join(root, file)
            rel_path = os.path.relpath(abs_path, base_dir)
            if is_source_file(file):
                res.append((abs_path, rel_path))
    return sorted(res, key=lambda x: x[1])


def add_code_paragraph(doc, code, font_name=CODE_FONT, font_size=CODE_FONT_SIZE):
    para = doc.add_paragraph()
    run = para.add_run(code)
    font = run.font
    font.name = font_name
    font.size = font_size
    para.style = "CodeStyle"


def main():
    files = walk_client_files(CLIENT_ROOT)
    doc = Document()
    style = doc.styles.add_style("CodeStyle", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = CODE_FONT
    style.font.size = Pt(9)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT)
    doc.add_heading("Исходники client (*.js, *.jsx, *.ts, *.tsx, *.css, *.scss)", 0)
    for abs_path, rel_path in files:
        doc.add_heading(rel_path, level=2)
        try:
            with open(abs_path, encoding="utf-8") as f:
                code = f.read()
        except Exception as e:
            code = f"[Ошибка чтения файла: {e}]"
        if code.strip():
            add_code_paragraph(doc, code)
    doc.save(OUT_FILENAME)
    print(f"Документ {OUT_FILENAME} создан! Файлов добавлено: {len(files)}")


if __name__ == "__main__":
    main()
