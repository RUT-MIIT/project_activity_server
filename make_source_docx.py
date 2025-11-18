import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt

EXCLUDE_PARTS = ["migrations", "__init__.py", "test", "tests", "venv"]
OUT_FILENAME = "project_source_code.docx"
CODE_FONT = "Consolas"
CODE_FONT_SIZE = Pt(9)


def is_source_file(file_path: str) -> bool:
    lower = file_path.lower()
    if any(part in lower for part in EXCLUDE_PARTS):
        return False
    return lower.endswith(".py")


def walk_py_files(base_dir):
    for root, dirs, files in os.walk(base_dir):
        # исключаем папку venv и прочие директории
        dirs[:] = [d for d in dirs if d not in EXCLUDE_PARTS]
        for f in files:
            path = os.path.join(root, f)
            rel = os.path.relpath(path, base_dir)
            if is_source_file(rel):
                yield rel


def add_code_paragraph(doc, code, font_name=CODE_FONT, font_size=CODE_FONT_SIZE):
    para = doc.add_paragraph()
    run = para.add_run(code)
    font = run.font
    font.name = font_name
    font.size = font_size
    para.style = "CodeStyle"


def main():
    base = os.path.dirname(os.path.abspath(__file__))
    files = list(walk_py_files(base))

    doc = Document()
    # Создаём стиль для кода
    style = doc.styles.add_style("CodeStyle", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = CODE_FONT
    style.font.size = Pt(9)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT)

    doc.add_heading("Исходный код Python (ключевые компоненты)", 0)
    for rel_path in sorted(files):
        doc.add_heading(rel_path, level=2)
        try:
            with open(rel_path, encoding="utf-8") as f:
                code = f.read()
        except Exception as e:
            code = f"[Ошибка чтения файла: {e}]"
        if code.strip():
            add_code_paragraph(doc, code)
    doc.save(OUT_FILENAME)
    print(f"Документ {OUT_FILENAME} создан! Файлов добавлено: {len(files)}")


if __name__ == "__main__":
    main()
